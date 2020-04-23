import webbrowser

from reportlab.pdfgen import canvas
from Generate_paper import selectModules
from fpdf import FPDF
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches



class printQuestion():
    def __init__(self):
        self.subject = ""
        self.q1 = []
        self.q1marks = ""
        self.q2 = []
        self.q2marks = ""
        self.q3 = []
        self.q3marks = ""
        self.q4 = []
        self.q4marks = ""
        self.q5 = []
        self.q5marks = ""
        self.q6 = []
        self.q6marks = ""
        self.date = ""
        self.marks = ""
        self.time =""
        self.sem = ""
        self.instructions = ""
        self.exam = ""
        self.fileName = ""

    def printQuestions(self):

        # q1scheme = self.check(self.q1marks)
        # print(q1scheme)
        # q2scheme = self.check(self.q2marks)
        # print(q2scheme)
        # q3scheme = self.check(self.q3marks)
        # print(q3scheme)
        # q4scheme = self.check(self.q4marks)
        # print(q4scheme)
        # q5scheme = self.check(self.q5marks)
        # print(q5scheme)
        # q6scheme = self.check(self.q6marks)
        # print(q6scheme)
        print("************************************")
        print(self.q1)
        print(self.q2)
        print(self.q3)
        print(self.q4)
        print(self.q5)
        print(self.q6)
        print(self.subject)
        document = Document()
        college = document.add_heading("SHIVAJIRAO S. JONDHALE COLLEGE OF ENGINEERING                DEPARTMENT OF IT ENGINEERING", level = 1)
        college.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sub = document.add_paragraph()
        sub.add_run(self.exam).bold = True
        sub.add_run("\nSubject : "+self.subject).bold = True
        sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        info = document.add_paragraph()
        info.add_run("                    "+self.sem+"                                                                                  Date : "+self.date).bold = True
        info.add_run("\n                     Time : "+self.time+"                                                                          Marks :"+ self.marks).bold = True
        sub.info = WD_PARAGRAPH_ALIGNMENT.CENTER
        inst =document.add_paragraph()
        paragraph_format = inst.paragraph_format
        paragraph_format.left_indent = Inches(0.63)
        if(self.instructions != ""):
            inst.add_run("Intructions :").bold = True
            inst.add_run("\n"+self.instructions)

        document.add_paragraph("________________________________________________________________________________________________________")
        if(len(self.q1) != 0):
            self.tm =0
            document.add_paragraph("")
            h1 =document.add_heading('Q1. Answer the following', level=2)
            for i in range(len(self.q1)):
                q1 = document.add_paragraph(str(i+1)+". "+self.q1[i][0])
                paragraph_format = q1.paragraph_format
                paragraph_format.left_indent = Inches(0.3)
                q1.add_run('   ('+str(self.q1[i][1])+' Marks)').bold = True
                self.tm += self.q1[i][1]
            h1.add_run("                                                                                  "+str(self.tm)+" Marks")

        if(len(self.q2) != 0):
            self.tm =0
            document.add_paragraph("")
            h2 =document.add_heading('Q2. Answer the following', level=2)
            for i in range(len(self.q2)):
                q2 = document.add_paragraph(str(i+1)+". "+self.q2[i][0])
                paragraph_format = q2.paragraph_format
                paragraph_format.left_indent = Inches(0.3)
                q2.add_run('   ('+str(self.q2[i][1])+' Marks)').bold = True
                self.tm += self.q2[i][1]
            h2.add_run("                                                                                  "+str(self.tm)+" Marks")

        if(len(self.q3) != 0):
            self.tm =0
            document.add_paragraph("")
            h3 =document.add_heading('Q3. Answer the following', level=2)
            for i in range(len(self.q3)):
                q3 =document.add_paragraph(str(i+1)+". "+self.q3[i][0])
                paragraph_format = q3.paragraph_format
                paragraph_format.left_indent = Inches(0.3)
                q3.add_run('   ('+str(self.q3[i][1])+' Marks)').bold = True
                self.tm += self.q3[i][1]
            h3.add_run("                                                                                  "+str(self.tm)+" Marks")

        if(len(self.q4) != 0):
            document.add_paragraph("")
            self.tm = 0
            h4 =document.add_heading('Q4. Answer the following', level=2)
            for i in range(len(self.q4)):
                q4 =document.add_paragraph(str(i+1)+". "+self.q4[i][0])
                paragraph_format = q4.paragraph_format
                paragraph_format.left_indent = Inches(0.3)
                q4.add_run('   ('+str(self.q4[i][1])+' Marks)').bold = True
                self.tm += self.q4[i][1]
            h4.add_run("                                                                                  "+str(self.tm)+" Marks")

        if(len(self.q5) != 0):
            self.tm =0
            document.add_paragraph("")
            h5 = document.add_heading('Q5. Answer the following', level=2)
            for i in range(len(self.q5)):
                q5 =document.add_paragraph(str(i+1)+". "+self.q5[i][0])
                paragraph_format = q5.paragraph_format
                paragraph_format.left_indent = Inches(0.3)
                q5.add_run('   ('+str(self.q5[i][1])+' Marks)').bold = True
                self.tm += self.q5[i][1]
            h5.add_run("                                                                                  "+str(self.tm)+" Marks")

        if(len(self.q6) != 0):
            self.tm = 0
            document.add_paragraph("")
            h6 =document.add_heading('Q6. Answer the following', level=2)
            for i in range(len(self.q6)):
                q6 =document.add_paragraph(str(i+1)+". "+self.q6[i][0])
                paragraph_format = q6.paragraph_format
                paragraph_format.left_indent = Inches(0.3)
                q6.add_run('   ('+str(self.q6[i][1])+' Marks)').bold = True
                self.tm += self.q6[i][1]
            h6.add_run("                                                                                  "+str(self.tm)+" Marks")



        document.save(self.fileName+'.docx')
        print("saved")

    def check(self, marks):
        if (marks == " 5, 5"):
            q = [5, 5]
        elif (marks == " 10"):
            q = [10]
        elif (marks == " 5, 5, 5"):
            q = [5, 5, 5]
        elif (marks == " 5, 10"):
            q = [5, 10]
        elif (marks == " 8, 8, 4"):
            q = [8, 8, 4]
        elif (marks == " 8, 6, 6"):
            q = [8, 6, 6]
        elif (marks == " 5, 5, 5, 5"):
            q = [5, 5, 5, 5]
        elif (marks == " 10, 10"):
            q = [10, 10]
        elif (marks == " 5, 5, 10"):
            q = [5, 5, 10]

        return q


